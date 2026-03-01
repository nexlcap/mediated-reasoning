import queue
import re
import tempfile
import threading
from pathlib import Path
from typing import List, Optional

import gradio as gr

from src.llm.client import ClaudeClient
from src.llm.prompts import build_pre_research_prompt
from src.mediator import Mediator
from src.project_memory import QAPair
from src.utils.document_loader import load_document, DocumentLoadError
from src.utils.formatters import (
    format_core_md,
    format_detail_md,
    format_sources_md,
)

ANSI_RE = re.compile(r'\033\[[0-9;]*m')

_START_RESEARCH_RE = re.compile(
    r'\b(start|begin|run|launch|kick\s*off|do)\b.{0,25}\bresearch\b', re.I
)

MODELS = [
    "claude-sonnet-4-6",
    "claude-opus-4-6",
    "claude-haiku-4-5-20251001",
    "gpt-4o",
    "gpt-4o-mini",
    "xai/grok-2",
    "xai/grok-3",
    "gemini/gemini-2.5-flash",
    "gemini/gemini-2.5-flash-lite",
    "gemini/gemini-2.5-pro",
]

# ── CSS ───────────────────────────────────────────────────────────────────────
_CSS = """
@import url('https://fonts.googleapis.com/css2?family=Inter:ital,wght@0,300;0,400;0,500;0,600;0,700;1,400&display=swap');

body,
.gradio-container,
.gradio-container p,
.gradio-container label,
.gradio-container input,
.gradio-container textarea,
.gradio-container select,
.gradio-container button,
.gradio-container h1,

.gradio-container h2,
.gradio-container h3,
.gradio-container h4 {
    font-family: 'Inter', ui-sans-serif, system-ui, -apple-system,
                 BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif !important;
}

/* ── App header ── */
.fusen-header { padding: 20px 0 8px 0; }
.fusen-title {
    font-size: 1.9em;
    font-weight: 700;
    background: linear-gradient(135deg, #f97316 0%, #ec4899 45%, #8b5cf6 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    display: inline-block;
    margin: 0 0 4px 0;
    line-height: 1.2;
}
.fusen-sub    { color: var(--body-text-color-subdued, #777); font-size: 0.9em;  margin: 2px 0 0 0; }

/* ── Analyze button ── */
.analyze-btn button {
    background: linear-gradient(135deg, #f97316 0%, #ec4899 100%) !important;
    border: none !important;
    color: white !important;
    font-weight: 600 !important;
    letter-spacing: 0.02em !important;
    transition: transform 0.15s ease, box-shadow 0.15s ease !important;
}
.analyze-btn button:hover {
    transform: translateY(-1px) !important;
    box-shadow: 0 6px 20px rgba(249, 115, 22, 0.35) !important;
}

/* ── Shared keyframes ── */
@keyframes statusFadeIn {
    from { opacity: 0; transform: translateY(8px); }
    to   { opacity: 1; transform: translateY(0);   }
}
@keyframes dotPulse {
    0%, 100% { opacity: 1;   }
    50%       { opacity: 0.2; }
}
@keyframes gradientShift {
    0%   { background-position: 0%   50%; }
    50%  { background-position: 100% 50%; }
    100% { background-position: 0%   50%; }
}

/* ── Pipeline animation ── */
.pipe-header {
    background: linear-gradient(270deg, #f97316, #ec4899, #8b5cf6, #3b82f6);
    background-size: 400% 400%;
    animation: gradientShift 3s ease infinite;
    color: white;
    padding: 10px 16px;
    border-radius: 8px;
    font-weight: 600;
    font-size: 0.88em;
    letter-spacing: 0.03em;
    margin-bottom: 14px;
}
.pipeline { padding: 4px 0; }
.pipe-step {
    display: flex;
    align-items: flex-start;
    gap: 10px;
    padding: 8px 4px;
    animation: statusFadeIn 0.4s ease both;
    border-bottom: 1px solid var(--border-color-primary, #e5e7eb);
}
.pipe-step:last-child { border-bottom: none; }
.pipe-dot {
    width: 9px; height: 9px;
    border-radius: 50%;
    margin-top: 4px;
    flex-shrink: 0;
    background: #d1d5db;
}
.pipe-done  .pipe-dot { background: #22c55e; }
.pipe-active .pipe-dot {
    background: #f97316;
    animation: dotPulse 1s ease infinite;
    box-shadow: 0 0 0 4px rgba(249, 115, 22, 0.18);
}
.pipe-text       { font-size: 0.85em; line-height: 1.45; color: var(--body-text-color, #333); }
.pipe-done .pipe-text { color: var(--body-text-color-subdued, #999); }
.status-error {
    color: #dc2626;
    padding: 8px 12px;
    border-left: 3px solid #dc2626;
    font-size: 0.9em;
}

/* ── Follow-up area ── */
.follow-qa-area { margin-top: 8px; }
.follow-qa-area p { line-height: 1.65; }

/* ── ChatGPT-style input pill ── */
.followup-row {
    margin-top: 12px !important;
    padding: 5px 6px 5px 4px !important;
    background: var(--background-fill-primary, #fff) !important;
    border: 1.5px solid var(--border-color-primary, #d9d9e3) !important;
    border-radius: 28px !important;
    align-items: center !important;
    gap: 2px !important;
    box-shadow: 0 2px 12px rgba(0,0,0,0.06) !important;
    transition: border-color 0.15s ease, box-shadow 0.15s ease !important;
}
.followup-row:focus-within {
    border-color: #f97316 !important;
    box-shadow: 0 0 0 3px rgba(249,115,22,0.10), 0 2px 12px rgba(0,0,0,0.06) !important;
}
/* Strip all borders/backgrounds from elements nested inside the pill */
.followup-row .block,
.followup-row .wrap,
.followup-row .container,
.followup-row label {
    border: none !important;
    box-shadow: none !important;
    background: transparent !important;
    padding: 0 !important;
    margin: 0 !important;
}
.followup-row textarea {
    border: none !important;
    outline: none !important;
    box-shadow: none !important;
    background: transparent !important;
    resize: none !important;
    padding: 6px 8px !important;
    font-size: 0.95em !important;
    line-height: 1.5 !important;
}

/* ── Resizable sidebars ──
   overflow is applied only when .open so the absolutely-positioned
   toggle button (left:100% when closed) is never clipped. ── */
#settings-sidebar.open {
    resize: horizontal !important;
    overflow: auto !important;
    min-width: 200px !important;
    max-width: 650px !important;
}
/* ── Main tabs — polished underline style ── */
#main-tabs {
    margin-top: 8px;
    border: none !important;
    background: transparent !important;
    box-shadow: none !important;
}
#main-tabs > .tab-nav {
    border-bottom: 2px solid var(--border-color-primary, #e5e7eb) !important;
    background: transparent !important;
    padding: 0 !important;
    gap: 4px !important;
    box-shadow: none !important;
}
#main-tabs > .tab-nav > button {
    font-size: 0.95em !important;
    font-weight: 500 !important;
    padding: 10px 22px !important;
    border: none !important;
    border-bottom: 2px solid transparent !important;
    margin-bottom: -2px !important;
    border-radius: 0 !important;
    background: transparent !important;
    color: var(--body-text-color-subdued, #888) !important;
    transition: color 0.15s ease, border-color 0.15s ease !important;
    letter-spacing: 0.01em !important;
    box-shadow: none !important;
}
#main-tabs > .tab-nav > button:hover {
    color: var(--body-text-color, #333) !important;
    background: transparent !important;
}
#main-tabs > .tab-nav > button.selected {
    color: #f97316 !important;
    border-bottom-color: #f97316 !important;
    font-weight: 600 !important;
    background: transparent !important;
}
#main-tabs > div:not(.tab-nav) {
    padding-top: 16px !important;
    border: none !important;
    box-shadow: none !important;
    background: transparent !important;
}


/* ── Memory buttons — match input field height & font ── */
.memory-btns button {
    font-size: 0.875rem !important;
    height: 40px !important;
    min-height: 40px !important;
}

/* ── 📎 upload button — ghost circle inside pill ── */
.chat-upload-btn button {
    border-radius: 50% !important;
    background: transparent !important;
    border: none !important;
    box-shadow: none !important;
    width: 34px !important;
    min-width: 34px !important;
    max-width: 34px !important;
    height: 34px !important;
    padding: 0 !important;
    color: var(--body-text-color-subdued, #888) !important;
    font-size: 1.05em !important;
    flex-shrink: 0 !important;
    transition: background 0.12s ease !important;
}
.chat-upload-btn button:hover {
    background: var(--background-fill-secondary, #f3f4f6) !important;
}

/* ── ↑ send button — solid dark circle inside pill ── */
#followup-btn button {
    border-radius: 50% !important;
    background: #1a1a1a !important;
    color: #fff !important;
    border: none !important;
    box-shadow: none !important;
    width: 34px !important;
    min-width: 34px !important;
    max-width: 34px !important;
    height: 34px !important;
    padding: 0 !important;
    font-size: 1.0em !important;
    flex-shrink: 0 !important;
    transition: background 0.12s ease !important;
}
#followup-btn button:hover {
    background: #333 !important;
    transform: none !important;
    box-shadow: none !important;
}

/* ── Research sidebar — section header weight ── */
#synthesis-section .label-wrap button,
#stats-section     .label-wrap button,
#breakdown-section .label-wrap button {
    font-weight: 600 !important;
}

/* ── Export ── */
.export-trigger-row {
    margin-top: 16px;
    justify-content: flex-end !important;
}
#export-btn button {
    font-size: 0.88em !important;
    font-weight: 500 !important;
    letter-spacing: 0.02em !important;
    border-radius: 8px !important;
    padding: 0 18px !important;
    height: 36px !important;
    min-height: 36px !important;
}
@keyframes exportFadeIn {
    from { opacity: 0; transform: translateY(-6px); }
    to   { opacity: 1; transform: translateY(0); }
}
#export-options {
    margin-top: 10px;
    padding: 20px 20px 16px !important;
    background: var(--background-fill-secondary, #f9fafb) !important;
    border: 1px solid var(--border-color-primary, #e5e7eb) !important;
    border-radius: 12px !important;
    box-shadow: 0 4px 18px rgba(0,0,0,0.07) !important;
    animation: exportFadeIn 0.16s ease both;
}
#export-options label {
    font-size: 0.82em !important;
    font-weight: 600 !important;
    text-transform: uppercase !important;
    letter-spacing: 0.055em !important;
    color: var(--body-text-color-subdued, #888) !important;
}
.export-actions { margin-top: 4px !important; }

/* ── Modals ── */
.modal-overlay > div.gr-group,
.modal-overlay {
    position: fixed !important;
    inset: 0 !important;
    background: rgba(0, 0, 0, 0.55) !important;
    z-index: 9999 !important;
    display: flex !important;
    align-items: center !important;
    justify-content: center !important;
    padding: 0 !important;
    margin: 0 !important;
    border: none !important;
    border-radius: 0 !important;
    box-shadow: none !important;
}
.modal-inner {
    background: var(--background-fill-primary, white) !important;
    border-radius: 12px !important;
    padding: 28px !important;
    width: 65vw !important;
    max-width: 900px !important;
    min-width: 380px !important;
    box-shadow: 0 24px 64px rgba(0, 0, 0, 0.4) !important;
    position: relative !important;
    inset: unset !important;
}
"""

# ── JavaScript: add title-attribute tooltips after Gradio renders ─────────────
_JS = """
function() {
    const TIPS = {
        "api-key":          "Your Anthropic / OpenAI / Google / xAI API key — never stored or shared",
        "memory-upload":     "Load a memory.md from a previous session to restore project context",
        "save-session":     "Download the current project memory as memory.md",
        "model-dd":         "LLM used for synthesis and the final report",
        "tavily":           "Optional Tavily premium search key — higher-quality citations than the default DuckDuckGo fallback",
        "analyze-btn":      "Synthesize your chat into a problem statement and run the full multi-specialist analysis",
        "show-qa":          "Load the full follow-up Q&A history into the right sidebar",
        "followup":         "Ask a follow-up question — the analysis context is kept so answers are grounded",
        "followup-btn":     "Send your follow-up question"
    };
    function applyTips() {
        for (const [id, tip] of Object.entries(TIPS)) {
            const el = document.getElementById(id);
            if (el) el.title = tip;
        }
    }
    applyTips();
    setTimeout(applyTips, 800);
    setTimeout(applyTips, 3000);

    // On mobile, close the settings sidebar so the main content is visible first
    function closeSidebarsOnMobile() {
        if (window.innerWidth >= 768) return;
        var el = document.getElementById('settings-sidebar');
        if (el && el.classList.contains('open')) {
            var btn = el.querySelector('button');
            if (btn) btn.click();
        }
    }
    setTimeout(closeSidebarsOnMobile, 500);
    setTimeout(closeSidebarsOnMobile, 1500);
}
"""

_HEADER_HTML = """
<div class="fusen-header">
  <div class="fusen-title">🔥 Fusen</div>
  <p class="fusen-sub">Experts debate, revise, and synthesize with live web search — business, career, research, strategy, and why the dinosaurs died out.</p>
</div>
"""


# ── Helpers ───────────────────────────────────────────────────────────────────
def _parse_weights(weights_str: str) -> dict:
    weights = {}
    for part in weights_str.split(","):
        part = part.strip()
        if not part:
            continue
        if "=" not in part:
            raise ValueError(f"Invalid weight '{part}' — expected format: agent=number")
        name, val = part.split("=", 1)
        weights[name.strip()] = float(val.strip())
    return weights



def _load_doc_context(document_paths):
    """Load documents from paths. Returns (DocumentContext | None, error_str | None)."""
    paths = document_paths if isinstance(document_paths, list) else ([document_paths] if document_paths else [])
    if not paths:
        return None, None
    try:
        docs = [load_document(p) for p in paths]
        if len(docs) == 1:
            return docs[0], None
        combined = "\n\n".join(
            f"## Document {i}: {d.filename}\n\n{d.content}"
            for i, d in enumerate(docs, 1)
        )
        return docs[0].__class__(
            filename=", ".join(d.filename for d in docs),
            content=combined,
            extraction_method=docs[0].extraction_method,
        ), None
    except DocumentLoadError as e:
        return None, str(e)


def _status_html(lines: List[str]) -> str:
    if not lines:
        return ""
    steps = "".join(
        f'<div class="pipe-step {"pipe-active" if i == len(lines)-1 else "pipe-done"}"'
        f' style="animation-delay:{min(i * 0.12, 2.0):.2f}s">'
        f'<div class="pipe-dot"></div>'
        f'<div class="pipe-text">{line}</div>'
        f'</div>'
        for i, line in enumerate(lines)
    )
    return (
        f'<div class="pipeline">'
        f'<div class="pipe-header">✦ Analysing your question…</div>'
        f'{steps}</div>'
    )


def _format_qa_right(qa_history: List[QAPair]) -> str:
    if not qa_history:
        return "*No follow-up questions yet.*"
    md = "### Follow-up history\n\n"
    for i, (q, a) in enumerate(qa_history, 1):
        md += f"**Q{i}.** {q}\n\n{a}\n\n---\n\n"
    return md


def _show_detail_in_sidebar(result) -> str:
    if result is None:
        return "*No analysis yet — run an analysis first.*"
    try:
        return format_detail_md(result, detailed=False)
    except Exception as e:
        return f"*Error loading detail: {e}*"



def _error_tuple(msg):
    """Return 9-tuple matching run_analysis error output."""
    return (f"<div class='status-error'>⚠ {msg}</div>",
            gr.update(), gr.update(), gr.update(),
            gr.update(), None, None, [],
            gr.update(selected="results-tab"))


def _extract_problem(history) -> str:
    """Derive a simple problem string from conversation history."""
    if not history:
        return ""
    return " ".join(q for q, _ in history if q)


def _synthesize_problem(history, model, api_key) -> str:
    """Ask the LLM to condense the pre-research conversation into a problem statement."""
    turns = []
    for q, a in (history or []):
        turns.append(f"User: {q}")
        if a:
            turns.append(f"Assistant: {a}")
    conversation = "\n".join(turns) or "(no conversation)"
    system = (
        "You are a concise problem distiller. Given a short pre-research conversation, "
        "write a single clear problem statement (2–4 sentences) suitable as input to a "
        "multi-agent analysis pipeline. Capture the core question, key constraints, and "
        "any relevant context. Output the problem statement only — no preamble."
    )
    user = f"Conversation:\n{conversation}\n\nProblem statement:"
    client = ClaudeClient(model=model, api_key=api_key.strip())
    return client.chat(system, user)


# ── Project helpers ───────────────────────────────────────────────────────────
def save_brief(brief_text: str):
    tmp = Path(tempfile.mkdtemp())
    brief_path = tmp / "memory.md"
    brief_path.write_text((brief_text or "").strip() + "\n", encoding="utf-8")
    return gr.update(visible=True, value=str(brief_path))


# ── Analysis ──────────────────────────────────────────────────────────────────
def run_analysis(
    problem, model, api_key,
    tavily_key,
    brief_text,
    document_paths=None,
):
    # 9 outputs: status_html · core_out · right_detail_md · right_stats_md ·
    #            followup_chatbot · result_state · mediator_state · qa_history_state · main_tabs

    def _progress(log_lines):
        return (_status_html(log_lines),
                gr.update(), gr.update(), gr.update(),  # core/detail/stats — no-op
                gr.update(),                             # chatbot — preserve pre-research msgs
                None, None, [],
                gr.update(selected="results-tab"))

    def _error(msg):
        return (f"<div class='status-error'>⚠ {msg}</div>",
                gr.update(), gr.update(), gr.update(),
                gr.update(), None, None, [],
                gr.update(selected="results-tab"))

    if not problem.strip():
        yield _error("Please enter a problem."); return
    if not api_key.strip():
        yield _error("Please enter your API key."); return

    document_context, doc_err = _load_doc_context(document_paths)
    if doc_err:
        yield _error(f"Document error: {doc_err}"); return

    effective_context = ""
    if brief_text and brief_text.strip():
        effective_context = f"Project memory:\n{brief_text.strip()}"

    key      = api_key.strip()
    status_q = queue.Queue()

    def on_progress(msg):
        status_q.put(("status", msg))

    client   = ClaudeClient(model=model, api_key=key)
    mediator = Mediator(
        client, search=True,
        deep_research=True, agent_client=None,
        repeat_prompt=True,
        tavily_api_key=(tavily_key or "").strip() or None,
        on_progress=on_progress,
        user_context=effective_context or None,
        document_context=document_context,
    )

    def run():
        try:
            status_q.put(("done", mediator.analyze(problem)))
        except BaseException as e:
            status_q.put(("error", str(e)))

    threading.Thread(target=run, daemon=True).start()

    log_lines = []
    while True:
        try:
            item = status_q.get(timeout=5)
        except Exception:
            if not threading.active_count():
                yield _error("Analysis thread exited unexpectedly."); return
            continue

        kind = item[0]
        if kind == "status":
            log_lines.append(item[1])
            yield _progress(log_lines)
        elif kind == "error":
            yield _error(item[1]); return
        elif kind == "done":
            result    = item[1]
            core_md   = format_core_md(result)
            detail_md = format_detail_md(result, detailed=False)
            q         = result.quality
            te        = {"good": "✅", "degraded": "⚠️", "poor": "❌"}.get(q.tier if q else "", "")
            stats     = ""
            if q:
                stats += f"{te} **Quality:** {q.tier} ({q.score:.2f})&ensp;"
            if result.timing:
                stats += f"⏱ **Time:** {result.timing.total_s:.0f}s&ensp;"
            if result.token_usage:
                stats += f"🔢 **Tokens:** {result.token_usage.total_input + result.token_usage.total_output:,}"
            sources_md = format_sources_md(result)
            right_top  = stats + ("\n\n---\n\n" + sources_md if sources_md else "")
            yield ("", core_md, detail_md, right_top,
                   [], result, mediator, [],
                   gr.update(selected="results-tab"))
            return


def _to_chatbot_msgs(qa_history: List[QAPair]) -> list:
    """Convert (question, answer) tuples to Gradio 6 message dicts."""
    msgs = []
    for q, a in qa_history:
        msgs.append({"role": "user",      "content": q})
        msgs.append({"role": "assistant", "content": a})
    return msgs


def run_followup(
    question, result, mediator, qa_history: List[QAPair],
    model, api_key, brief_text, document_paths,
):
    # Outputs: followup_chatbot · qa_history_state · followup_input · analysis_trigger_state
    history = list(qa_history or [])
    q = (question or "").strip()
    if not q:
        yield _to_chatbot_msgs(history), history, "", False
        return

    if result is None:
        # ── Pre-research phase ────────────────────────────────────────────────
        if not api_key or not api_key.strip():
            yield (
                _to_chatbot_msgs(history)
                + [{"role": "user",      "content": q},
                   {"role": "assistant", "content": "⚠ Please enter your API key in the settings panel before chatting."}],
                history, "", False,
            )
            return

        # Phrase detection — auto-trigger analysis
        if _START_RESEARCH_RE.search(q):
            ack_msgs = (
                _to_chatbot_msgs(history)
                + [{"role": "user",      "content": q},
                   {"role": "assistant", "content": "Starting the analysis now! 🔬 One moment…"}]
            )
            new_history = history + [(q, "Starting the analysis now! 🔬 One moment…")]
            yield ack_msgs, new_history, "", True
            return

        doc_ctx, doc_err = _load_doc_context(document_paths)
        if doc_err:
            yield (
                _to_chatbot_msgs(history)
                + [{"role": "user",      "content": q},
                   {"role": "assistant", "content": f"⚠ Document error: {doc_err}"}],
                history, "", False,
            )
            return

        doc_text = doc_ctx.content if doc_ctx else None
        effective_context = (brief_text or "").strip() or None
        problem_str = _extract_problem(history)
        system, messages = build_pre_research_prompt(
            problem_str, doc_text, effective_context, q, history
        )
        client = ClaudeClient(model=model, api_key=api_key.strip())

        # Show user message immediately
        yield _to_chatbot_msgs(history) + [{"role": "user", "content": q}], history, "", False

        full = ""
        for chunk in client.chat_stream(system, messages):
            full += ANSI_RE.sub('', chunk)
            yield (
                _to_chatbot_msgs(history)
                + [{"role": "user",      "content": q},
                   {"role": "assistant", "content": full}],
                history, "", False,
            )

        final = history + [(q, full)]
        yield _to_chatbot_msgs(final), final, "", False

    else:
        # ── Post-research phase ───────────────────────────────────────────────
        if mediator is None:
            yield _to_chatbot_msgs(history), history, "", False
            return

        # Show user message immediately; assistant slot will fill via streaming
        yield _to_chatbot_msgs(history) + [{"role": "user", "content": q}], history, "", False

        full = ""
        for chunk in mediator.followup_stream(result, q, history):
            full += ANSI_RE.sub('', chunk)
            yield (
                _to_chatbot_msgs(history)
                + [{"role": "user",      "content": q},
                   {"role": "assistant", "content": full}],
                history, "", False,
            )

        final = history + [(q, full)]
        yield _to_chatbot_msgs(final), final, "", False


def maybe_start_analysis(trigger, history, model, api_key, tavily_key, brief_text, doc_paths):
    """Fired after run_followup via .then(); only acts when trigger=True."""
    # 10 outputs: 9 from run_analysis + analysis_trigger_state
    _no_op = tuple(gr.update() for _ in range(9)) + (False,)
    if not trigger:
        yield _no_op
        return
    if not api_key or not api_key.strip():
        yield _no_op
        return
    problem = _synthesize_problem(history, model, api_key)
    for update in run_analysis(problem, model, api_key, tavily_key, brief_text, doc_paths):
        yield update + (False,)


def _start_research(history, model, api_key, tavily_key, brief_text, doc_paths):
    """Fired by 'Start Research' button click."""
    if not history:
        yield _error_tuple("Please describe your problem in the chat before starting research.")
        return
    if not api_key or not api_key.strip():
        yield _error_tuple("Please enter your API key.")
        return
    problem = _synthesize_problem(history, model, api_key)
    yield from run_analysis(problem, model, api_key, tavily_key, brief_text, doc_paths)


def _handle_chat_upload(new_files, existing_paths, chatbot_msgs):
    """Handle 📎 UploadButton — append file paths to documents_state and ack in chat."""
    paths = list(existing_paths or [])
    if isinstance(new_files, list):
        paths.extend(new_files)
    else:
        paths.append(new_files)
    uploaded = new_files if isinstance(new_files, list) else [new_files]
    names = ", ".join(Path(p).name for p in uploaded if p)
    msg = {"role": "assistant", "content": f"📎 Attached: **{names}**"}
    return paths, (chatbot_msgs or []) + [msg]


# ── Export ────────────────────────────────────────────────────────────────────
_ANSI_RE2 = re.compile(r'\033\[[0-9;]*m')
_HTML_ENT_RE = re.compile(r'&ensp;|&nbsp;')
_INLINE_MD_RE = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')


def _clean_md(text: str) -> str:
    """Strip ANSI codes and normalise HTML entities."""
    text = _ANSI_RE2.sub('', text)
    text = _HTML_ENT_RE.sub(' ', text)
    text = re.sub(r'&amp;', '&', text)
    text = re.sub(r'&lt;', '<', text)
    text = re.sub(r'&gt;', '>', text)
    return text


def _md_to_plain(text: str) -> str:
    """Reduce markdown to readable plain text."""
    text = re.sub(r'^\s*#{1,6}\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = re.sub(r'^\s*>\s+', '  ', text, flags=re.MULTILINE)
    text = re.sub(r'^\s*[-*]\s+', '• ', text, flags=re.MULTILINE)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    text = re.sub(r'^---+$', '', text, flags=re.MULTILINE)
    return text


def _build_export_parts(result, sections: list) -> list:
    """Return [(title, markdown_text)] for each selected section."""
    if result is None:
        return []
    parts = []
    if "Synthesis" in sections:
        md = format_core_md(result)
        if md.strip():
            parts.append(("Synthesis", md))
    if "Stats & Sources" in sections:
        q = result.quality
        te = {"good": "✅", "degraded": "⚠️", "poor": "❌"}.get(q.tier if q else "", "")
        lines = []
        if q:
            lines.append(f"**Quality:** {te} {q.tier} (score: {q.score:.2f})")
        if result.timing:
            lines.append(f"**Total time:** {result.timing.total_s:.0f}s")
        if result.token_usage:
            total_tok = result.token_usage.total_input + result.token_usage.total_output
            lines.append(f"**Tokens used:** {total_tok:,}")
        sources_md = format_sources_md(result)
        content = "\n\n".join(lines)
        if sources_md:
            content += "\n\n---\n\n" + sources_md
        if content.strip():
            parts.append(("Stats & Sources", content))
    if "Agent Breakdown" in sections:
        md = format_detail_md(result, detailed=False)
        if md.strip():
            parts.append(("Agent Breakdown", md))
    return parts


# ── Text export ──
def _export_as_text(parts: list) -> gr.update:
    lines = ["# Fusen Analysis Report", ""]
    for title, md in parts:
        lines += [f"## {title}", "", _clean_md(md), "", "---", ""]
    tmp = Path(tempfile.mktemp(suffix=".md"))
    tmp.write_text("\n".join(lines), encoding="utf-8")
    return gr.update(visible=True, value=str(tmp))


# ── Word export ──
def _add_inline_runs(para, text: str):
    """Write **bold**, *italic*, `code` spans as Word runs."""
    for chunk in _INLINE_MD_RE.split(text):
        if chunk.startswith('**') and chunk.endswith('**') and len(chunk) > 4:
            para.add_run(chunk[2:-2]).bold = True
        elif chunk.startswith('*') and chunk.endswith('*') and len(chunk) > 2:
            para.add_run(chunk[1:-1]).italic = True
        elif chunk.startswith('`') and chunk.endswith('`') and len(chunk) > 2:
            run = para.add_run(chunk[1:-1])
            run.font.name = 'Courier New'
        elif chunk:
            para.add_run(chunk)


def _md_body_to_docx(doc, md: str):
    for line in md.splitlines():
        s = line.strip()
        if not s:
            continue
        if re.match(r'^---+$', s):
            doc.add_paragraph()
            continue
        m = re.match(r'^(#{1,4})\s+(.*)', s)
        if m:
            doc.add_heading(m.group(2), min(len(m.group(1)), 4))
            continue
        m = re.match(r'^>\s+(.*)', s)
        if m:
            p = doc.add_paragraph(style='Quote')
            _add_inline_runs(p, m.group(1))
            continue
        m = re.match(r'^[-*]\s+(.*)', s)
        if m:
            p = doc.add_paragraph(style='List Bullet')
            _add_inline_runs(p, m.group(1))
            continue
        p = doc.add_paragraph()
        _add_inline_runs(p, s)


def _export_as_docx(parts: list) -> gr.update:
    from docx import Document
    from docx.shared import RGBColor
    doc = Document()
    heading = doc.add_heading("Fusen Analysis Report", 0)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0xF9, 0x73, 0x16)
    for title, md in parts:
        doc.add_heading(title, 1)
        _md_body_to_docx(doc, _clean_md(md))
    tmp = Path(tempfile.mktemp(suffix=".docx"))
    doc.save(str(tmp))
    return gr.update(visible=True, value=str(tmp))


# ── PDF export ──
_PDF_FONT_PAIRS = [
    ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
     "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
    ("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
     "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf"),
    ("/usr/share/fonts/truetype/freefont/FreeSans.ttf",
     "/usr/share/fonts/truetype/freefont/FreeSansBold.ttf"),
    ("/System/Library/Fonts/Supplemental/Arial.ttf",
     "/System/Library/Fonts/Supplemental/Arial Bold.ttf"),
    ("C:/Windows/Fonts/arial.ttf", "C:/Windows/Fonts/arialbd.ttf"),
]


def _export_as_pdf(parts: list) -> gr.update:
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_left_margin(20)
    pdf.set_right_margin(20)
    pdf.set_x(20)

    fn = "Helvetica"
    for reg, bold in _PDF_FONT_PAIRS:
        if Path(reg).exists():
            pdf.add_font("Body", fname=reg)
            pdf.add_font("Body", style="B",
                         fname=bold if Path(bold).exists() else reg)
            fn = "Body"
            break

    def safe(t: str) -> str:
        if fn == "Helvetica":
            return t.encode("latin-1", errors="replace").decode("latin-1")
        return t

    def mc(text, h=5, **kw):
        """multi_cell with cursor always reset to l_margin."""
        pdf.multi_cell(0, h, text,
                       new_x=XPos.LMARGIN, new_y=YPos.NEXT, **kw)

    # Document title
    pdf.set_font(fn, style="B", size=18)
    pdf.set_text_color(249, 115, 22)
    mc(safe("Fusen Analysis Report"), h=9)
    pdf.set_text_color(30, 30, 30)
    pdf.ln(5)

    for title, md in parts:
        # Section heading bar
        pdf.set_font(fn, style="B", size=12)
        pdf.set_fill_color(243, 244, 246)
        pdf.set_text_color(50, 50, 70)
        mc(safe(title), h=8, fill=True)
        pdf.set_text_color(30, 30, 30)
        pdf.ln(3)

        # Body — strip markdown, render line by line
        plain = _md_to_plain(_clean_md(md))
        pdf.set_font(fn, size=10)
        for line in plain.splitlines():
            s = line.strip()
            if not s:
                pdf.ln(2)
                continue
            mc(safe(s))
        pdf.ln(6)

    tmp = Path(tempfile.mktemp(suffix=".pdf"))
    pdf.output(str(tmp))
    return gr.update(visible=True, value=str(tmp))


def _export_analysis(result, sections, fmt):
    if result is None or not sections:
        return gr.update(visible=False)
    parts = _build_export_parts(result, sections)
    if not parts:
        return gr.update(visible=False)
    if fmt == "Text":
        return _export_as_text(parts)
    if fmt == "Word":
        return _export_as_docx(parts)
    return _export_as_pdf(parts)


# ── UI ────────────────────────────────────────────────────────────────────────
_HEAD_META = """
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">

<!-- Primary SEO -->
<title>Fusen — Multi-Agent AI Analysis for Business, Career &amp; Strategy</title>
<meta name="description" content="Fusen runs a panel of AI specialists that independently debate, revise, and synthesize answers with live web search. Ask any complex question — business strategy, career decisions, research, investments — and get a structured, conflict-aware report in minutes.">
<meta name="keywords" content="multi-agent AI, AI analysis, decision support AI, AI brainstorming, business strategy AI, AI research assistant, expert AI panel, multi-perspective analysis, AI debate, AI synthesis, Fusen AI">
<meta name="author" content="Fusen">
<meta name="robots" content="index, follow">
<link rel="canonical" href="https://app.fusen.ai/">

<!-- Open Graph (Facebook, LinkedIn, etc.) -->
<meta property="og:type" content="website">
<meta property="og:url" content="https://app.fusen.ai/">
<meta property="og:title" content="Fusen — Multi-Agent AI Analysis for Business, Career &amp; Strategy">
<meta property="og:description" content="A panel of AI specialists independently debate, revise, and synthesize answers with live web search. Get structured, conflict-aware analysis for any complex question — business, career, research, or strategy.">
<meta property="og:site_name" content="Fusen">
<meta property="og:locale" content="en_US">

<!-- Twitter Card -->
<meta name="twitter:card" content="summary_large_image">
<meta name="twitter:title" content="Fusen — Multi-Agent AI Analysis">
<meta name="twitter:description" content="AI specialists debate and synthesize expert answers with live web search. Business, career, research, strategy.">
<meta name="twitter:site" content="@fusenai">

<!-- JSON-LD structured data -->
<script type="application/ld+json">
{
  "@context": "https://schema.org",
  "@type": "SoftwareApplication",
  "name": "Fusen",
  "url": "https://app.fusen.ai/",
  "applicationCategory": "BusinessApplication",
  "operatingSystem": "Web",
  "description": "Fusen uses multi-agent mediated reasoning to analyze complex problems from multiple expert perspectives. Specialists independently analyze, revise in light of each other's views, and produce a synthesized report with live web citations.",
  "offers": {
    "@type": "Offer",
    "price": "0",
    "priceCurrency": "USD"
  },
  "featureList": [
    "Multi-agent AI panel with dynamically selected specialist roles",
    "Live web search with DuckDuckGo and optional Tavily integration",
    "Three-round reasoning: independent analysis, cross-agent revision, synthesis",
    "Structured conflict extraction and resolution",
    "Follow-up Q&A grounded in the full analysis context",
    "Support for PDF, DOCX, XLSX, and other document uploads",
    "Project memory across sessions",
    "Multiple LLM backends: Claude, GPT-4o, Gemini, Grok"
  ]
}
</script>
"""

with gr.Blocks(title="Fusen — Multi-Agent AI Analysis") as demo:
    result_state           = gr.State(None)
    mediator_state         = gr.State(None)
    qa_history_state       = gr.State([])
    documents_state        = gr.State([])
    analysis_trigger_state = gr.State(False)
    export_panel_state     = gr.State(False)

    # ── Left Sidebar ──────────────────────────────────────────────────────────
    with gr.Sidebar(label="Settings", open=True, position="left", width=320,
                    elem_id="settings-sidebar"):

        api_key_input = gr.Textbox(
            label="API key", placeholder="sk-ant-… · sk-… · xai-…",
            type="password", elem_id="api-key",
        )
        model_dd = gr.Dropdown(
            choices=MODELS, value=MODELS[0], label="Model", elem_id="model-dd",
        )

        tavily_input = gr.Textbox(
            label="Tavily API key\n(optional)",
            placeholder="tvly-…", type="password", elem_id="tavily",
        )

        with gr.Row(elem_classes=["memory-btns"]):
            brief_upload = gr.UploadButton(
                "Load memory", file_types=[".md", ".txt"],
                type="filepath", elem_id="memory-upload",
            )
            save_memory_btn = gr.Button("Save memory", variant="secondary",
                                        elem_id="save-session")
        brief_load_info = gr.Markdown(visible=False)
        brief_download  = gr.File(show_label=False, visible=False,
                                  elem_id="memory-download")
        brief_area = gr.Textbox(visible=False, elem_id="memory")


    # ── Main canvas ───────────────────────────────────────────────────────────
    gr.HTML(_HEADER_HTML)

    with gr.Tabs(elem_id="main-tabs") as main_tabs:

        with gr.TabItem("💬 Chat", id="chat-tab"):
            followup_chatbot = gr.Chatbot(
                show_label=False,
                height=400,
                elem_classes=["follow-qa-area"],
                layout="bubble",
                buttons=["copy"],
                placeholder="Describe your problem here — ask questions, upload documents, and chat before clicking Start Research. Or just type your question and hit Start Research.",
            )

            with gr.Row(elem_classes=["followup-row"]):
                chat_upload_btn = gr.UploadButton(
                    "📎",
                    file_types=[".pdf", ".txt", ".md", ".rst", ".docx", ".pptx", ".xlsx", ".xls"],
                    file_count="multiple",
                    type="filepath",
                    scale=1,
                    elem_classes=["chat-upload-btn"],
                    elem_id="chat-upload",
                )
                followup_input = gr.Textbox(
                    show_label=False, placeholder="Ask a question or describe your problem…",
                    lines=1, max_lines=6, scale=8, elem_id="followup",
                )
                followup_btn = gr.Button("↑", scale=1, variant="secondary",
                                         elem_id="followup-btn")

            submit_btn = gr.Button("Start Research", variant="primary", size="lg",
                                   elem_classes=["analyze-btn"], elem_id="analyze-btn")

        with gr.TabItem("📊 Results", id="results-tab"):
            status_html = gr.HTML()

            with gr.Accordion("📋 Synthesis", open=True, elem_id="synthesis-section"):
                core_out = gr.Markdown("*Run an analysis to see results here.*")

            with gr.Accordion("📊 Stats & Sources", open=True, elem_id="stats-section"):
                right_stats_md = gr.Markdown()

            with gr.Accordion("🔍 Agent Breakdown", open=False, elem_id="breakdown-section"):
                right_detail_md = gr.Markdown("*Agent detail appears here after analysis.*")

            with gr.Row(elem_classes=["export-trigger-row"]):
                export_btn = gr.Button(
                    "Export ↓", variant="secondary", scale=0,
                    min_width=120, elem_id="export-btn",
                )

            with gr.Column(visible=False, elem_id="export-options") as export_options:
                with gr.Row():
                    export_sections = gr.CheckboxGroup(
                        choices=["Synthesis", "Stats & Sources", "Agent Breakdown"],
                        value=["Synthesis", "Stats & Sources", "Agent Breakdown"],
                        label="Sections",
                        scale=3,
                    )
                    export_format = gr.Radio(
                        choices=["PDF", "Word", "Text"],
                        value="PDF",
                        label="Format",
                        scale=1,
                    )
                with gr.Row(elem_classes=["export-actions"]):
                    export_cancel_btn = gr.Button("Cancel", scale=1, variant="secondary")
                    export_download_btn = gr.Button("Download", scale=2, variant="primary")
                export_file = gr.File(show_label=False, visible=False,
                                      elem_id="export-file")

    # ── Wiring ────────────────────────────────────────────────────────────────
    def _load_brief(p):
        if not p:
            return "", gr.update(visible=False)
        return (
            Path(p).read_text(encoding="utf-8"),
            gr.update(visible=True, value=f"📄 {Path(p).name}"),
        )

    brief_upload.upload(
        fn=_load_brief, inputs=[brief_upload], outputs=[brief_area, brief_load_info],
    )

    chat_upload_btn.upload(
        fn=_handle_chat_upload,
        inputs=[chat_upload_btn, documents_state, followup_chatbot],
        outputs=[documents_state, followup_chatbot],
    )

    submit_btn.click(
        fn=_start_research,
        inputs=[
            qa_history_state,
            model_dd,
            api_key_input,
            tavily_input,
            brief_area,
            documents_state,
        ],
        outputs=[
            status_html, core_out, right_detail_md, right_stats_md,
            followup_chatbot, result_state, mediator_state, qa_history_state,
            main_tabs,
        ],
    )

    _fu_in = [
        followup_input, result_state, mediator_state, qa_history_state,
        model_dd, api_key_input, brief_area, documents_state,
    ]
    _fu_out = [followup_chatbot, qa_history_state, followup_input, analysis_trigger_state]

    _then_in = [
        analysis_trigger_state, qa_history_state, model_dd, api_key_input,
        tavily_input, brief_area, documents_state,
    ]
    _then_out = [
        status_html, core_out, right_detail_md, right_stats_md,
        followup_chatbot, result_state, mediator_state, qa_history_state,
        main_tabs, analysis_trigger_state,
    ]

    followup_btn.click(fn=run_followup, inputs=_fu_in, outputs=_fu_out)\
        .then(fn=maybe_start_analysis, inputs=_then_in, outputs=_then_out)
    followup_input.submit(fn=run_followup, inputs=_fu_in, outputs=_fu_out)\
        .then(fn=maybe_start_analysis, inputs=_then_in, outputs=_then_out)

    save_memory_btn.click(
        fn=save_brief, inputs=[brief_area], outputs=[brief_download],
    )

    def _toggle_export(is_open):
        new = not is_open
        label = "Export ↑" if new else "Export ↓"
        return gr.update(visible=new), new, gr.update(value=label)

    def _close_export():
        return gr.update(visible=False), False, gr.update(value="Export ↓")

    export_btn.click(
        fn=_toggle_export,
        inputs=[export_panel_state],
        outputs=[export_options, export_panel_state, export_btn],
    )
    export_cancel_btn.click(
        fn=_close_export,
        outputs=[export_options, export_panel_state, export_btn],
    )
    export_download_btn.click(
        fn=_export_analysis,
        inputs=[result_state, export_sections, export_format],
        outputs=[export_file],
    )

demo.launch(server_name="0.0.0.0", server_port=7860, ssr_mode=False,
            theme=gr.themes.Soft(), css=_CSS, js=_JS, head=_HEAD_META)
